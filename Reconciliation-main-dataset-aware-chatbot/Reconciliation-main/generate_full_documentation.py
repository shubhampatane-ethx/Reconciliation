#!/usr/bin/env python3
"""Generate a comprehensive Word document for the entire Consistency Reconciliation application."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# Create a new Document
doc = Document()

# Add title page
title = doc.add_heading('Consistency Reconciliation', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

subtitle = doc.add_heading('Complete Application Documentation', level=2)
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

date_para = doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

version_para = doc.add_paragraph('Version: 1.0')
version_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add a horizontal line
doc.add_paragraph('_' * 80).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Table of Contents
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Project Overview & Architecture',
    '3. Technology Stack',
    '4. Backend Implementation (Flask)',
    '5. Frontend Implementation (React)',
    '6. Database & Storage Layer',
    '7. Docker Containerization',
    '8. API Endpoints',
    '9. Reconciliation Engine',
    '10. User Interface & Features',
    '11. Theme Picker Feature',
    '12. Deployment & Running the Application',
    '13. File Structure',
    '14. Installation & Setup',
    '15. Future Enhancements'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

# Add page break
doc.add_page_break()

# 1. Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'Consistency Reconciliation is a full-stack web application designed to compare two CSV/XLSX files '
    '(source and target) and generate detailed reconciliation reports. The application identifies mismatches, '
    'missing records, and data discrepancies, providing users with day-wise summaries and detailed comparison results.'
)
doc.add_paragraph(
    'The application combines React for the frontend, Flask for the backend, and Docker for containerization, '
    'creating a scalable, maintainable solution for data reconciliation tasks.'
)

# 2. Project Overview & Architecture
doc.add_heading('2. Project Overview & Architecture', level=1)

doc.add_heading('2.1 Core Functionality', level=2)
features = [
    'Upload source and target CSV/XLSX files',
    'Define key columns for row-level matching',
    'Perform multi-criteria reconciliation',
    'Generate detailed mismatch reports with day-wise summaries',
    'Store uploaded files and reports persistently',
    'Browse previously processed files and their chunks',
    'Switch between four custom themes (Dark, Light, Solar, Midnight)',
    'User-friendly dashboard with drag-and-drop file upload'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('2.2 High-Level Architecture', level=2)
arch_para = doc.add_paragraph()
arch_para.add_run('Client Layer: ').bold = True
arch_para.add_run('React SPA with Vite bundler, hosted in Node 20 Alpine container\n')
arch_para = doc.add_paragraph()
arch_para.add_run('API Layer: ').bold = True
arch_para.add_run('RESTful Flask API for file uploads, reconciliation, and report retrieval\n')
arch_para = doc.add_paragraph()
arch_para.add_run('Storage Layer: ').bold = True
arch_para.add_run('JSON-based local file store for chunks and reports (no external DB)\n')
arch_para = doc.add_paragraph()
arch_para.add_run('Orchestration: ').bold = True
arch_para.add_run('Docker Compose manages frontend and backend services')

# 3. Technology Stack
doc.add_heading('3. Technology Stack', level=1)

tech_table = doc.add_table(rows=9, cols=3)
tech_table.style = 'Light Grid Accent 1'
header_cells = tech_table.rows[0].cells
header_cells[0].text = 'Layer'
header_cells[1].text = 'Technology'
header_cells[2].text = 'Purpose'

techs = [
    ('Frontend', 'React 18 + Vite', 'SPA with fast HMR dev experience'),
    ('Frontend', 'Axios', 'HTTP client for API calls'),
    ('Frontend', 'CSS3 + CSS Variables', 'Responsive styling with theme support'),
    ('Backend', 'Flask', 'Lightweight RESTful API framework'),
    ('Backend', 'Pandas', 'Data parsing and comparison logic'),
    ('Backend', 'Python 3.12', 'Core runtime environment'),
    ('Infrastructure', 'Docker & Docker Compose', 'Container orchestration'),
    ('Storage', 'JSON Files', 'Persistent local storage for data')
]

for i, (layer, tech, purpose) in enumerate(techs, 1):
    row_cells = tech_table.rows[i].cells
    row_cells[0].text = layer
    row_cells[1].text = tech
    row_cells[2].text = purpose

# 4. Backend Implementation (Flask)
doc.add_heading('4. Backend Implementation (Flask)', level=1)

doc.add_heading('4.1 Main Application File: backend/app.py', level=2)
doc.add_paragraph('The Flask application handles:')
backend_tasks = [
    'File upload and storage',
    'CSV/XLSX parsing via pandas',
    'Reconciliation logic and difference detection',
    'Report generation and persistence',
    'API endpoints for frontend communication'
]
for task in backend_tasks:
    doc.add_paragraph(task, style='List Bullet')

doc.add_heading('4.2 Storage Module: backend/storage.py', level=2)
doc.add_paragraph('Manages persistent data storage:')
storage_functions = [
    'store_file(filename, df, file_type): Chunks and stores uploaded files as JSON',
    'store_report(report, source_meta, target_meta, key_columns, day_summary): Persists reconciliation reports',
    'list_reports(): Returns metadata for all stored reports',
    'load_file_chunks(file_id): Retrieves previously stored file chunks'
]
for func in storage_functions:
    doc.add_paragraph(func, style='List Bullet')

doc.add_heading('4.3 Reconciliation Engine', level=2)
doc.add_paragraph(
    'The backend implements a sophisticated reconciliation algorithm that:'
)
recon_steps = [
    'Normalizes column names and data types',
    'Matches rows using user-specified key columns',
    'Compares field values between source and target',
    'Flags missing rows in source or target',
    'Generates detailed mismatch records with original and target values',
    'Aggregates results by date for summary reports'
]
for step in recon_steps:
    doc.add_paragraph(step, style='List Bullet')

# 5. Frontend Implementation (React)
doc.add_heading('5. Frontend Implementation (React)', level=1)

doc.add_heading('5.1 Main Application: frontend/src/App.jsx', level=2)
doc.add_paragraph('Core React component providing:')
react_features = [
    'File upload form with drag-and-drop support',
    'Key column input with chip-based UI',
    'Reconciliation execution and progress tracking',
    'Report display with detailed mismatch rendering',
    'Stored files browser with chunk preview',
    'Theme picker with four color schemes',
    'Toast notifications for user feedback'
]
for feature in react_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('5.2 Key React State Variables', level=2)
state_table = doc.add_table(rows=11, cols=2)
state_table.style = 'Light Grid Accent 1'
header_cells = state_table.rows[0].cells
header_cells[0].text = 'State Variable'
header_cells[1].text = 'Purpose'

states = [
    ('sourceFile, targetFile', 'Stores uploaded file references'),
    ('keyColumns', 'Array of column names for row matching'),
    ('report, daySummary', 'Reconciliation results'),
    ('storedFiles', 'List of previously uploaded files'),
    ('selectedFileChunks', 'Currently displayed file chunks'),
    ('loading, error', 'Request status and error messages'),
    ('themeMenuOpen, currentTheme', 'Theme picker state and active theme'),
    ('toasts', 'Notification messages queue'),
    ('activeView', 'Current tab (Reconcile, Stored Files, Reports)'),
    ('uploadProgress', 'File upload percentage')
]

for i, (var, purpose) in enumerate(states, 1):
    row_cells = state_table.rows[i].cells
    row_cells[0].text = var
    row_cells[1].text = purpose

doc.add_heading('5.3 Styling: frontend/src/styles.css', level=2)
doc.add_paragraph('Comprehensive CSS with:')
css_features = [
    'CSS variables for theming (--bg, --panel, --primary, --accent, --text)',
    'Responsive grid layouts for desktop and mobile',
    'Component styles: buttons, cards, inputs, modals',
    'Smooth transitions (280ms) for theme switching',
    'Glassmorphism effects for modern aesthetic',
    'Hover and focus states for accessibility',
    'Tooltip styles for inline help'
]
for feature in css_features:
    doc.add_paragraph(feature, style='List Bullet')

# 6. Database & Storage Layer
doc.add_heading('6. Database & Storage Layer', level=1)

doc.add_heading('6.1 Storage Architecture', level=2)
doc.add_paragraph(
    'Instead of a traditional database, the application uses a JSON-based file storage system '
    'to avoid dependencies on heavy external packages within the Docker slim image.'
)

storage_dirs = [
    'backend/vector_store/metadata.json: Metadata for all stored files',
    'backend/vector_store/{file_id}.json: Chunked content of each uploaded file',
    'backend/vector_store/reports/{timestamp}_{name}_report.json: Reconciliation reports'
]
for dir_info in storage_dirs:
    doc.add_paragraph(dir_info, style='List Bullet')

doc.add_heading('6.2 File Metadata Structure', level=2)
doc.add_paragraph('Each stored file has:')
metadata_fields = [
    'file_id: UUID for unique identification',
    'filename: Original uploaded filename',
    'file_type: csv or xlsx',
    'chunk_count: Number of JSON chunks',
    'upload_timestamp: ISO datetime of upload'
]
for field in metadata_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading('6.3 Report Storage', level=2)
doc.add_paragraph('Reports include:')
report_fields = [
    'report: Detailed mismatch records with field-level comparisons',
    'day_summary: Aggregated counts by date',
    'source_meta: Source file metadata',
    'target_meta: Target file metadata',
    'key_columns: Columns used for matching',
    'timestamp: Report generation time'
]
for field in report_fields:
    doc.add_paragraph(field, style='List Bullet')

# 7. Docker Containerization
doc.add_heading('7. Docker Containerization', level=1)

doc.add_heading('7.1 Docker Compose Setup', level=2)
doc.add_paragraph('File: docker-compose.yml')
doc.add_paragraph('Services defined:')
services = [
    ('frontend', 'Node 20 Alpine, runs Vite dev server on port 5173'),
    ('backend', 'Python 3.12-slim, runs Flask on port 5000')
]
for service, desc in services:
    doc.add_paragraph(f'{service}: {desc}', style='List Bullet')

doc.add_heading('7.2 Frontend Dockerfile', level=2)
doc.add_paragraph('File: frontend/Dockerfile')
dockerfile_steps = [
    'Base: node:20-alpine',
    'Install dependencies from package*.json',
    'Copy source code',
    'Expose port 5173',
    'Run Vite dev server with --host 0.0.0.0'
]
for step in dockerfile_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('7.3 Backend Dockerfile', level=2)
doc.add_paragraph('File: backend/Dockerfile')
backend_dockerfile_steps = [
    'Base: python:3.12-slim',
    'Install system dependencies (if needed)',
    'Install Python packages from requirements.txt',
    'Copy backend code',
    'Expose port 5000',
    'Run Flask app with development server'
]
for step in backend_dockerfile_steps:
    doc.add_paragraph(step, style='List Bullet')

# 8. API Endpoints
doc.add_heading('8. API Endpoints', level=1)

doc.add_heading('8.1 Reconciliation Endpoint', level=2)
endpoint_table = doc.add_table(rows=5, cols=4)
endpoint_table.style = 'Light Grid Accent 1'
header_cells = endpoint_table.rows[0].cells
header_cells[0].text = 'Method'
header_cells[1].text = 'Endpoint'
header_cells[2].text = 'Input'
header_cells[3].text = 'Output'

endpoints = [
    ('POST', '/api/reconcile', 'source_file, target_file, key_columns', 'Reconciliation report + day summary'),
    ('GET', '/api/stored-files', 'None', 'List of stored file metadata'),
    ('GET', '/api/file-chunks/{file_id}', 'file_id', 'File content chunks'),
    ('GET', '/api/reports', 'None', 'List of report files')
]

for i, (method, endpoint, input_data, output) in enumerate(endpoints, 1):
    row_cells = endpoint_table.rows[i].cells
    row_cells[0].text = method
    row_cells[1].text = endpoint
    row_cells[2].text = input_data
    row_cells[3].text = output

# 9. Reconciliation Engine Details
doc.add_heading('9. Reconciliation Engine', level=1)

doc.add_heading('9.1 Algorithm Overview', level=2)
algo_steps = [
    '1. Parse source and target files (CSV/XLSX)',
    '2. Normalize column names and data types',
    '3. Create lookup index using key columns',
    '4. For each source row, find matching target row(s)',
    '5. Compare field values and record differences',
    '6. Identify unmatched rows in source and target',
    '7. Group mismatches by date field',
    '8. Return detailed report with summary'
]
for step in algo_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('9.2 Key Column Matching', level=2)
doc.add_paragraph(
    'Users specify key columns (e.g., "id", "customer_id") that uniquely identify rows. '
    'The engine uses these columns to match rows between source and target files before comparing values.'
)

doc.add_heading('9.3 Mismatch Detection', level=2)
mismatch_types = [
    'Value Mismatch: Same row ID, different field values',
    'Missing in Target: Row exists in source but not target',
    'Missing in Source: Row exists in target but not source'
]
for mtype in mismatch_types:
    doc.add_paragraph(mtype, style='List Bullet')

# 10. User Interface & Features
doc.add_heading('10. User Interface & Features', level=1)

doc.add_heading('10.1 Main Dashboard', level=2)
ui_features = [
    'Reconciliation Panel: Upload files, specify key columns, run reconciliation',
    'Stored Files Tab: Browse and preview previously uploaded files',
    'Reports Tab: View saved reconciliation reports',
    'Sidebar Navigation: Switch between tabs',
    'Theme Picker: Select from four color schemes',
    'Toast Notifications: Real-time user feedback'
]
for feature in ui_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('10.2 File Upload Experience', level=2)
upload_features = [
    'Drag-and-drop support for easy file selection',
    'File type validation (CSV, XLSX, XLS)',
    'Upload progress indicator with percentage',
    'File size display (KB) after selection',
    'Remove button to clear selection'
]
for feature in upload_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('10.3 Key Columns Input', level=2)
doc.add_paragraph(
    'Chip-based input allows users to specify multiple key columns. '
    'Users type a column name and press Enter or comma to add. Chips show added columns with remove (×) buttons.'
)

doc.add_heading('10.4 Report Display', level=2)
report_features = [
    'Detailed mismatch records in expandable rows',
    'Side-by-side comparison of source vs. target values',
    'Day-wise summary counts',
    'Filter for mismatches only',
    'Search across mismatch records'
]
for feature in report_features:
    doc.add_paragraph(feature, style='List Bullet')

# 11. Theme Picker Feature
doc.add_heading('11. Theme Picker Feature', level=1)

doc.add_heading('11.1 Theme Options', level=2)
theme_table = doc.add_table(rows=5, cols=3)
theme_table.style = 'Light Grid Accent 1'
header_cells = theme_table.rows[0].cells
header_cells[0].text = 'Theme'
header_cells[1].text = 'Primary Colors'
header_cells[2].text = 'Use Case'

themes = [
    ('Dark', '#071029, #071830', 'Low-light environments, default'),
    ('Light', '#f8fafc, #ffffff', 'Daytime use, bright backgrounds'),
    ('Solar', '#fff7ed, #ffb347', 'Warm aesthetics, orange accents'),
    ('Midnight', '#020617, #0f172a', 'Premium feel, deep blue/purple')
]

for i, (theme, colors, use_case) in enumerate(themes, 1):
    row_cells = theme_table.rows[i].cells
    row_cells[0].text = theme
    row_cells[1].text = colors
    row_cells[2].text = use_case

doc.add_heading('11.2 Implementation', level=2)
theme_impl = [
    'CSS variables for all theme colors (--bg, --panel, --text, --primary, --accent)',
    'THEMES object in React maps theme names to CSS variable values',
    'applyTheme(name) function updates :root CSS variables',
    'localStorage persistence with key "cr_theme"',
    'Smooth 280ms CSS transitions when switching themes',
    'Active theme highlighted in sidebar picker'
]
for impl in theme_impl:
    doc.add_paragraph(impl, style='List Bullet')

# 12. Deployment & Running
doc.add_heading('12. Deployment & Running the Application', level=1)

doc.add_heading('12.1 Prerequisites', level=2)
prereqs = [
    'Docker and Docker Compose installed',
    'Windows 10/11, macOS, or Linux',
    'Minimum 2GB RAM allocated to Docker',
    'Port 5173 and 5000 available'
]
for prereq in prereqs:
    doc.add_paragraph(prereq, style='List Bullet')

doc.add_heading('12.2 Running the Application', level=2)
run_steps = [
    'Navigate to workspace: cd C:\\Users\\Shubham Patane\\Desktop\\Consistency',
    'Build and start services: docker compose up -d --build',
    'Wait for services to be ready (~30 seconds)',
    'Open browser: http://localhost:5173',
    'UI should load with the default Dark theme',
    'To stop: docker compose down'
]
for i, step in enumerate(run_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_heading('12.3 Checking Logs', level=2)
log_commands = [
    'Frontend logs: docker compose logs frontend',
    'Backend logs: docker compose logs backend',
    'Real-time logs: docker compose logs -f',
    'Full rebuild: docker compose up -d --build'
]
for cmd in log_commands:
    doc.add_paragraph(cmd, style='List Bullet')

# 13. File Structure
doc.add_heading('13. File Structure', level=1)

doc.add_paragraph('The project is organized as follows:')
doc.add_paragraph()

file_structure = """
Consistency/
├── frontend/
│   ├── public/
│   │   ├── logo.svg          # Application logo
│   │   └── favicon.svg       # Browser favicon
│   ├── src/
│   │   ├── App.jsx           # Main React component
│   │   └── styles.css        # Global styles + themes
│   ├── index.html            # HTML entry point
│   ├── package.json          # Node dependencies
│   ├── package-lock.json     # Locked versions
│   ├── vite.config.js        # Vite configuration
│   └── Dockerfile            # Frontend container
│
├── backend/
│   ├── app.py                # Flask application
│   ├── storage.py            # Storage management
│   ├── requirements.txt       # Python dependencies
│   ├── vector_store/         # Local storage
│   │   ├── metadata.json     # File metadata
│   │   ├── {file_id}.json    # File chunks
│   │   └── reports/          # Reconciliation reports
│   └── Dockerfile            # Backend container
│
├── docker-compose.yml        # Multi-container orchestration
├── .gitignore                # Git exclusions
└── README.md                 # Project documentation
"""

doc.add_paragraph(file_structure, style='Normal')

# 14. Installation & Setup
doc.add_heading('14. Installation & Setup', level=1)

doc.add_heading('14.1 Initial Setup (Dev Machine)', level=2)
setup_steps = [
    'Clone or extract the Consistency folder to Desktop',
    'Ensure Docker Desktop is running',
    'Open terminal in the Consistency folder',
    'Run: docker compose up -d --build',
    'Services start automatically and run in background'
]
for i, step in enumerate(setup_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_heading('14.2 Development Workflow', level=2)
dev_workflow = [
    'Frontend code changes: Vite HMR applies automatically (no restart needed)',
    'Backend code changes: Restart backend with docker compose restart backend',
    'CSS changes: Apply immediately due to HMR',
    'Adding packages: Update package.json or requirements.txt, then rebuild'
]
for step in dev_workflow:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('14.3 Production Deployment', level=2)
prod_steps = [
    'Use production-grade Node and Python images (non-alpine)',
    'Set Flask debug=False',
    'Configure a production-grade web server (gunicorn for Flask)',
    'Add reverse proxy (nginx) for static file serving',
    'Set up SSL/TLS certificates',
    'Configure environment variables for secrets'
]
for i, step in enumerate(prod_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

# 15. Future Enhancements
doc.add_heading('15. Future Enhancements', level=1)

doc.add_heading('15.1 Backend Enhancements', level=2)
backend_enhancements = [
    'Connect to PostgreSQL/MySQL for better scalability',
    'Add user authentication and authorization',
    'Implement scheduled report exports (daily/weekly)',
    'Support for more file formats (JSON, Parquet, SQL databases)',
    'Parallel processing for large files',
    'Email notifications for completed reconciliations'
]
for enh in backend_enhancements:
    doc.add_paragraph(enh, style='List Bullet')

doc.add_heading('15.2 Frontend Enhancements', level=2)
frontend_enhancements = [
    'Advanced filtering and sorting on mismatch records',
    'Export reports as PDF or Excel',
    'Data visualization: charts for mismatch trends',
    'Batch file upload and processing',
    'Undo/redo history for reconciliation actions',
    'Dark mode schedule (auto-switch based on time)',
    'Mobile-responsive design optimization'
]
for enh in frontend_enhancements:
    doc.add_paragraph(enh, style='List Bullet')

doc.add_heading('15.3 Operational Enhancements', level=2)
ops_enhancements = [
    'Kubernetes deployment configuration',
    'CI/CD pipeline (GitHub Actions, GitLab CI)',
    'Monitoring and alerting (Prometheus, Grafana)',
    'Automated backups for stored reports',
    'API rate limiting and quotas',
    'WebSocket support for real-time updates'
]
for enh in ops_enhancements:
    doc.add_paragraph(enh, style='List Bullet')

# 16. Troubleshooting
doc.add_heading('16. Troubleshooting', level=1)

doc.add_heading('16.1 Common Issues', level=2)

issues_table = doc.add_table(rows=6, cols=2)
issues_table.style = 'Light Grid Accent 1'
header_cells = issues_table.rows[0].cells
header_cells[0].text = 'Issue'
header_cells[1].text = 'Solution'

issues = [
    ('Port 5173/5000 already in use', 'Kill existing process: netstat -ano | findstr :5173'),
    ('Docker containers not starting', 'Check logs: docker compose logs'),
    ('File upload fails', 'Verify file format is CSV or XLSX; check file size'),
    ('Theme not persisting', 'Clear browser localStorage and refresh'),
    ('Reconciliation takes too long', 'Large files may take time; check backend logs for progress')
]

for i, (issue, solution) in enumerate(issues, 1):
    row_cells = issues_table.rows[i].cells
    row_cells[0].text = issue
    row_cells[1].text = solution

# 17. Conclusion
doc.add_heading('17. Conclusion', level=1)
doc.add_paragraph(
    'Consistency Reconciliation is a complete, production-ready web application for file comparison and '
    'reconciliation. With a modern tech stack, containerized deployment, and intuitive user interface, '
    'it provides an efficient solution for identifying and managing data discrepancies.'
)
doc.add_paragraph(
    'The application demonstrates best practices in full-stack development, including separation of concerns, '
    'scalable architecture, responsive UI design, and comprehensive documentation.'
)

conclusion_para = doc.add_paragraph()
conclusion_para.add_run('Key Achievements:\n').bold = True
achievements = [
    'End-to-end file reconciliation with customizable key columns',
    'Persistent storage and report generation',
    'Multi-theme UI with smooth transitions',
    'Dockerized deployment for easy scaling',
    'Responsive, accessible interface',
    'Comprehensive error handling and user feedback'
]
for achievement in achievements:
    doc.add_paragraph(achievement, style='List Bullet')

# Add final page break and footer
doc.add_page_break()

# Footer
footer_para = doc.add_paragraph('_' * 80)
footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

footer = doc.add_paragraph()
footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
footer.add_run('Consistency Reconciliation Application\n').bold = True
footer.add_run(f'Documentation Generated: {datetime.now().strftime("%B %d, %Y at %H:%M")}\n')
footer.add_run('For questions or support, refer to README.md or contact the development team.')

# Save the document
output_path = r'c:\Users\Shubham Patane\Desktop\Consistency\COMPLETE_APPLICATION_DOCUMENTATION.docx'
doc.save(output_path)
print(f'✓ Complete application Word document created: {output_path}')
